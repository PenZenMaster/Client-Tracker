r"""
Module/Script Name: chatgpt_background.py
Path: E:\projects\Project Tracking\chatgpt_background.py

Description:
Generates business background summaries using OpenAI GPT-4. Scrapes website content
and creates comprehensive service and background information documents for SEO clients.

Author(s):
Rank Rocket Co (C) Copyright 2025 - All Rights Reserved

Created Date:
2024-01-15

Last Modified Date:
2025-10-23

Version:
v1.02

License:
CC BY-SA 4.0 - https://creativecommons.org/licenses/by-sa/4.0/

Comments:
* v1.02 - Added type hints and Google-style docstrings
* v1.01 - Added standardized file header and ASCII-only output
* v1.00 - Initial release with GPT-4 background generation
"""

import os
from typing import Dict, Any
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document  # type: ignore[import-not-found]
from scraper import scrape_website_text

# Load API Key from .env
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def run(client_config: Dict[str, Any]) -> None:
    """Generate business background summary document using GPT-4.

    Scrapes website content and uses OpenAI GPT-4 to create a comprehensive
    business summary document for SEO client onboarding and content strategy.

    Workflow:
    1. Scrapes text from client website URL
    2. Sends scraped content to GPT-4 with business context
    3. Generates professional background summary
    4. Saves summary to DOCX file in client output directory

    Args:
        client_config: Configuration dictionary containing:
            - name (str): Business name
            - address (str, optional): Business address. Defaults to "Unknown"
            - url (str, optional): Website URL to scrape. Defaults to ""
            - output_root (str): Base output directory path

    Returns:
        None. Writes DOCX file to: {output_root}/{name}/{name} background information.docx

    Example:
        >>> config = {
        ...     "name": "ABC Heating & Cooling",
        ...     "address": "123 Main St, Phoenix, AZ 85001",
        ...     "url": "https://www.abchvac.com",
        ...     "output_root": "./output"
        ... }
        >>> run(config)
        Generating background summary for ABC Heating & Cooling...
        Saved background info to ./output/ABC Heating & Cooling/ABC Heating & Cooling background information.docx
    """
    client_name = client_config["name"]
    address = client_config.get("address", "Unknown")
    url = client_config.get("url", "").strip()
    base_output = client_config["output_root"]

    if not base_output.lower().endswith(client_name.lower()):
        output_path = os.path.join(base_output, client_name)
    else:
        output_path = base_output

    os.makedirs(output_path, exist_ok=True)

    # Scrape site content
    site_text = scrape_website_text(url)

    # Use triple quotes to avoid escape hell
    prompt = f"""Based on the following website content, provide a summary of services
and background information for this HVAC company.

Business Name: {client_name}
Address: {address}
Website Content:
{site_text}
"""

    print(f"Generating background summary for {client_name}...")

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an SEO assistant."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
    )

    summary = response.choices[0].message.content

    doc = Document()
    doc.add_heading(f"{client_name} – Background Information", 0)
    doc.add_paragraph(summary)
    output_file = os.path.join(
        output_path, f"{client_name} background information.docx"
    )
    doc.save(output_file)

    print(f"Saved background info to {output_file}")
